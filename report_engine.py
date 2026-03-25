# -*- coding: utf-8 -*-
"""
report_engine.py — 通用报告生成引擎

用法：
    # 1. 扫描模板，生成 CSV 配置骨架（傻瓜模式第一步）
    python report_engine.py --scan templates/my_report.docx

    # 2. 编辑生成的 templates/my_report.csv，填写 SQL / 固定值

    # 3. 生成报告
    python report_engine.py --template templates/my_report.docx
    python report_engine.py          # 批量处理模板目录

模板目录约定（--template-dir 指定，默认 ./templates）：
    templates/
        my_report.docx   ← Word 模板，用 {{变量名}} 标记占位符
        my_report.csv    ← 字段配置（新，与模板同名，扩展名换 .csv）
        my_report.json   ← 字段配置（旧，依然兼容）

也可用 --template / --config 分别指定路径。

配置文件格式见 templates/README.md。
"""

import argparse
import csv
import datetime
import json
import os
import re
import sys

import pymysql
from docx import Document
from docx.oxml.ns import qn

# ── 环境变量支持 ──────────────────────────────────────────────
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

DB_CONFIG = {
    'host':     os.getenv('DB_HOST',     'localhost'),
    'port':     int(os.getenv('DB_PORT', '3306')),
    'user':     os.getenv('DB_USER',     'root'),
    'password': os.getenv('DB_PASSWORD', ''),
    'database': os.getenv('DB_NAME',     'daily'),
    'charset':  'utf8mb4',
}


# ── 内置日期变量 ──────────────────────────────────────────────
def _builtin_dates():
    today     = datetime.date.today()
    yesterday = today - datetime.timedelta(days=1)
    last_week = yesterday - datetime.timedelta(days=7)
    return {
        'today':          str(today),
        'yesterday':      str(yesterday),
        'last_week':      str(last_week),
        'date':           datetime.datetime.now().strftime('%Y%m%d_%H%M%S'),
        # 中文格式（直接在模板中使用，无需 CSV 配置）
        'yesterday_md':  f'{yesterday.month}月{yesterday.day}日',
        'yesterday_ymd': f'{yesterday.year}年{yesterday.month}月{yesterday.day}日',
        'today_md':      f'{today.month}月{today.day}日',
        'today_ymd':     f'{today.year}年{today.month}月{today.day}日',
    }


# ── 格式化函数 ────────────────────────────────────────────────
def _fmt(value, fmt: str) -> str:
    if value is None:
        return ''
    try:
        v = float(value)
    except (TypeError, ValueError):
        return str(value)

    if fmt == 'wan':
        return f'{v / 10000:.2f}万'
    if fmt == 'pct':
        sign = '+' if v >= 0 else ''
        return f'{sign}{v * 100:.2f}%'
    if fmt == 'int':
        return f'{int(v):,}'
    if fmt == 'date_cn':
        # 支持 date 对象或 YYYY-MM-DD 字符串
        try:
            d = datetime.date.fromisoformat(str(value)[:10])
            return f'{d.month}月{d.day}日'
        except Exception:
            return str(value)
    return str(value)


# ── SQL 查询 ──────────────────────────────────────────────────
def _run_sql(query: str, conn) -> object:
    """执行 SQL，返回第一行第一列的值；无结果返回 None。"""
    with conn.cursor() as cur:
        cur.execute(query)
        row = cur.fetchone()
        return row[0] if row else None


# ── 变量解析 ──────────────────────────────────────────────────
def resolve_variables(config: dict, conn) -> dict:
    """
    根据配置解析所有变量，返回 {变量名: 格式化后的字符串}。
    支持 __AUTO_DISCOVER__ 标记（CSV 模式下只填字段名时自动找表）。
    """
    dates = _builtin_dates()
    result = dict(dates)

    db_name = DB_CONFIG.get('database', 'daily')

    for name, spec in config.get('variables', {}).items():
        vtype = spec.get('type', 'raw')
        fmt   = spec.get('format', 'raw')

        if vtype == 'date':
            val = spec.get('value', 'today')
            raw = dates.get(val, val)
        elif vtype == 'sql':
            raw_query = spec['query']
            # 自动发现表名
            if raw_query.startswith('__AUTO_DISCOVER__:'):
                field = raw_query.split(':', 1)[1]
                try:
                    raw_query = _resolve_auto_discover(field, conn, db_name)
                except ValueError as e:
                    print(f'  [WARN] {e}')
                    result[name] = f'{{{{{name}}}}}'
                    continue
            # 将内置日期变量替换进 SQL
            for k, v in dates.items():
                raw_query = raw_query.replace(f'{{{k}}}', v)
            try:
                raw = _run_sql(raw_query, conn)
            except Exception as e:
                print(f'  [WARN] 变量 {name} 查询失败: {e}')
                raw = None
        elif vtype == 'literal':
            raw = spec.get('value', '')
        else:
            raw = spec.get('value', '')

        result[name] = _fmt(raw, fmt)

    return result


# ── docx 占位符扫描 ───────────────────────────────────────────
_PLACEHOLDER_RE = re.compile(r'\{\{(\w+)\}\}')

# 引擎内置的日期变量（无需用户配置 SQL）
_BUILTIN_DATE_KEYS = {
    'today', 'yesterday', 'last_week', 'date',
    'yesterday_md', 'yesterday_ymd', 'today_md', 'today_ymd',
}


def scan_placeholders(docx_path: str) -> list:
    """从 docx 模板中提取所有 {{变量名}} 占位符，去重后按字母排序。"""
    doc = Document(docx_path)
    found = set()

    def _collect(para):
        for m in _PLACEHOLDER_RE.finditer(para.text):
            found.add(m.group(1))

    for para in doc.paragraphs:
        _collect(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _collect(para)
    for section in doc.sections:
        for hdr in (section.header, section.footer):
            if hdr:
                for para in hdr.paragraphs:
                    _collect(para)

    return sorted(found)


def generate_csv_scaffold(docx_path: str, csv_path: str = None) -> str:
    """
    扫描 docx 模板，在同目录生成 CSV 配置骨架（业务人员友好版）。

    生成的 CSV 列：
      占位符      模板中的 {{变量名}}
      数据库字段  填写 表名.字段名（如 platform_daily_metrics.platform_dau）
                  或只填字段名（系统自动查找所在表）
      格式        万 / 百分比 / 整数 / 日期 / 原始
      固定值      若填写则直接使用该文字，不查数据库
      备注        仅供人工阅读
    """
    placeholders = scan_placeholders(docx_path)
    if not placeholders:
        print(f'[WARN] 模板 {docx_path} 中未发现任何 {{{{变量}}}} 占位符')
        return ''

    if csv_path is None:
        csv_path = os.path.splitext(docx_path)[0] + '.csv'

    # 若 CSV 已存在，保留用户已填写的内容
    existing = {}
    if os.path.exists(csv_path):
        with open(csv_path, encoding='utf-8-sig', newline='') as f:
            for row in csv.DictReader(f):
                name = row.get('占位符', '').strip()
                if name:
                    existing[name] = row

    fieldnames = ['占位符', '数据库字段', '格式', '固定值', '备注']
    rows = []
    for name in placeholders:
        if name in existing:
            rows.append(existing[name])
        elif name in _BUILTIN_DATE_KEYS:
            rows.append({
                '占位符': name,
                '数据库字段': '',
                '格式': '日期',
                '固定值': '',
                '备注': '内置日期变量，无需填写',
            })
        else:
            rows.append({
                '占位符': name,
                '数据库字段': '',
                '格式': '原始',
                '固定值': '',
                '备注': '请填写数据库字段（如 表名.字段名）或在固定值列填写文字',
            })

    with open(csv_path, 'w', newline='', encoding='utf-8-sig') as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)

    print(f'[OK] CSV 配置骨架已生成：{csv_path}')
    print(f'     共 {len(placeholders)} 个占位符：{", ".join(placeholders)}')
    print(f'     用 Excel 打开 CSV，在"数据库字段"列填写 表名.字段名，保存后运行引擎')
    return csv_path


# 格式名称中英文映射
_FMT_ALIAS = {
    '万': 'wan', '万人次': 'wan', '万次': 'wan',
    '百分比': 'pct', '%': 'pct', '百分': 'pct',
    '整数': 'int', '数字': 'int',
    '日期': 'date_cn', '中文日期': 'date_cn',
    '原始': 'raw', '不格式化': 'raw', '': 'raw',
}


def _normalize_fmt(fmt_str: str) -> str:
    """将中文格式名转为引擎内部格式键。"""
    s = (fmt_str or '').strip()
    return _FMT_ALIAS.get(s, s) or 'raw'


def _field_to_sql(field: str) -> str:
    """
    将 '表名.字段名' 或 '字段名' 转为 SELECT 语句。
    若字段值以 SELECT 开头，视为完整 SQL 直接透传（供开发者预填复杂计算）。
    日期过滤默认取昨日（stat_date = '{yesterday}'）。
    """
    field = field.strip()
    if field.upper().startswith('SELECT'):
        return field  # 完整 SQL，直接透传
    if '.' in field:
        table, col = field.split('.', 1)
        return f"SELECT {col.strip()} FROM {table.strip()} WHERE stat_date = '{{yesterday}}'"
    # 只有字段名，运行时通过 INFORMATION_SCHEMA 自动发现表
    return f"__AUTO_DISCOVER__:{field}"


def _resolve_auto_discover(field: str, conn, db_name: str) -> str:
    """查询 INFORMATION_SCHEMA 找到字段所在的表，返回完整 SQL。"""
    with conn.cursor() as cur:
        cur.execute(
            "SELECT TABLE_NAME FROM INFORMATION_SCHEMA.COLUMNS "
            "WHERE TABLE_SCHEMA = %s AND COLUMN_NAME = %s LIMIT 1",
            (db_name, field)
        )
        row = cur.fetchone()
    if not row:
        raise ValueError(f'数据库中未找到字段 "{field}"，请在 CSV 中改为 表名.字段名 格式')
    table = row[0]
    print(f'  [INFO] 字段 {field} 自动发现所在表：{table}')
    return f"SELECT {field} FROM {table} WHERE stat_date = '{{yesterday}}'"


def csv_to_config(csv_path: str) -> dict:
    """
    将业务人员填写的 CSV 解析为引擎内部 config dict。

    CSV 列：
      占位符      对应模板 {{变量名}}
      数据库字段  表名.字段名 或 字段名（自动找表）；留空则用固定值或内置日期
      格式        万/百分比/整数/日期/原始（支持中英文）
      固定值      直接填入文字，不查数据库
      备注        忽略
    """
    variables = {}

    with open(csv_path, encoding='utf-8-sig', newline='') as f:
        reader = csv.DictReader(f)
        for row in reader:
            name = (row.get('占位符') or '').strip()
            if not name or name.startswith('#'):
                continue

            db_field  = (row.get('数据库字段') or '').strip()
            fmt_raw   = (row.get('格式') or '').strip()
            fixed_val = (row.get('固定值') or '').strip()
            fmt       = _normalize_fmt(fmt_raw)

            if fixed_val:
                # 固定值优先
                variables[name] = {'type': 'literal', 'value': fixed_val, 'format': fmt}
            elif not db_field and name in _BUILTIN_DATE_KEYS:
                # 内置日期变量
                variables[name] = {'type': 'date', 'value': name, 'format': fmt or 'date_cn'}
            elif db_field:
                sql = _field_to_sql(db_field)
                variables[name] = {'type': 'sql', 'query': sql, 'format': fmt}
            else:
                print(f'  [WARN] 占位符 "{name}" 未填写数据库字段也无固定值，将保留原占位符')

    base_name = os.path.splitext(os.path.basename(csv_path))[0]
    return {
        'report_name':     base_name,
        'output_dir':      '报表产出',
        'output_filename': f'{base_name}_{{date}}.docx',
        'variables':       variables,
    }


# ── Word 模板替换 ─────────────────────────────────────────────


def _replace_in_run(run, variables: dict):
    text = run.text
    def replacer(m):
        key = m.group(1)
        return variables.get(key, m.group(0))
    run.text = _PLACEHOLDER_RE.sub(replacer, text)


def _replace_in_paragraph(para, variables: dict):
    # 先尝试逐 run 替换（保留格式）
    for run in para.runs:
        _replace_in_run(run, variables)

    # 如果占位符跨 run 被拆散，做全文合并替换（会丢失部分格式，但保证内容正确）
    full = para.text
    if _PLACEHOLDER_RE.search(full):
        def replacer(m):
            return variables.get(m.group(1), m.group(0))
        new_text = _PLACEHOLDER_RE.sub(replacer, full)
        if new_text != full:
            # 清空所有 run，写入第一个 run
            for i, run in enumerate(para.runs):
                run.text = new_text if i == 0 else ''


def fill_template(template_path: str, variables: dict, output_path: str):
    doc = Document(template_path)

    # 正文段落
    for para in doc.paragraphs:
        _replace_in_paragraph(para, variables)

    # 表格单元格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, variables)

    # 页眉页脚
    for section in doc.sections:
        for hdr in (section.header, section.footer):
            if hdr:
                for para in hdr.paragraphs:
                    _replace_in_paragraph(para, variables)

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f'[OK] 报告已生成：{output_path}')


# ── 主流程 ────────────────────────────────────────────────────
def generate(template_path: str, config_path: str):
    # 支持 .csv 和 .json 两种配置格式
    if config_path.lower().endswith('.csv'):
        config = csv_to_config(config_path)
    else:
        with open(config_path, encoding='utf-8') as f:
            config = json.load(f)

    report_name  = config.get('report_name', '报告')
    output_dir   = config.get('output_dir', '报表产出')
    filename_tpl = config.get('output_filename', f'{report_name}_{{date}}.docx')

    print(f'正在生成：{report_name}')
    print(f'  模板：{template_path}')
    print(f'  配置：{config_path}')

    conn = pymysql.connect(**DB_CONFIG)
    try:
        variables = resolve_variables(config, conn)
    finally:
        conn.close()

    filename = filename_tpl
    for k, v in variables.items():
        filename = filename.replace(f'{{{k}}}', v)

    output_path = os.path.join(output_dir, filename)
    fill_template(template_path, variables, output_path)
    return output_path


# ── CLI ───────────────────────────────────────────────────────
def _find_config(template_path: str) -> str:
    """同名配置文件自动发现：优先 .csv，其次 .json。"""
    base = os.path.splitext(template_path)[0]
    for ext in ('.csv', '.CSV', '.json', '.JSON'):
        p = base + ext
        if os.path.exists(p):
            return p
    return ''


def main():
    parser = argparse.ArgumentParser(description='通用报告生成引擎')
    parser.add_argument('--template',     help='Word 模板路径（.docx）')
    parser.add_argument('--config',       help='字段配置路径（.csv 或 .json）')
    parser.add_argument('--scan',         metavar='DOCX',
                        help='扫描 docx 模板，生成 CSV 配置骨架（傻瓜模式第一步）')
    parser.add_argument('--template-dir', default='templates',
                        help='模板目录（默认 ./templates），自动扫描所有 .docx+配置 对')
    parser.add_argument('--list',         action='store_true',
                        help='列出模板目录中可用的报告模板')
    args = parser.parse_args()

    # ── --scan：生成 CSV 骨架 ──────────────────────────────────
    if args.scan:
        generate_csv_scaffold(args.scan)
        return

    # ── --list：列出模板 ───────────────────────────────────────
    if args.list:
        tdir = args.template_dir
        if not os.path.isdir(tdir):
            print(f'模板目录不存在：{tdir}')
            sys.exit(1)
        pairs = []
        for f in sorted(os.listdir(tdir)):
            if not f.lower().endswith('.docx'):
                continue
            cfg = _find_config(os.path.join(tdir, f))
            if cfg:
                cfg_type = '(CSV)' if cfg.lower().endswith('.csv') else '(JSON)'
                pairs.append((f, f'[OK] {cfg_type}'))
            else:
                pairs.append((f, '[--] (缺少 .csv 或 .json 配置)'))
        if not pairs:
            print(f'模板目录 {tdir} 中没有找到 .docx 文件')
        else:
            print(f'模板目录：{tdir}')
            for name, status in pairs:
                print(f'  {status}  {name}')
        return

    # ── 生成报告 ───────────────────────────────────────────────
    if args.template:
        template_path = args.template
        config_path   = args.config or _find_config(template_path)
        if not config_path:
            print(f'未找到配置文件，请先运行：python report_engine.py --scan {template_path}')
            sys.exit(1)
        generate(template_path, config_path)
    else:
        tdir = args.template_dir
        if not os.path.isdir(tdir):
            print(f'模板目录不存在：{tdir}，请用 --template 指定单个模板')
            sys.exit(1)
        found = False
        for f in sorted(os.listdir(tdir)):
            if not f.lower().endswith('.docx'):
                continue
            tpl = os.path.join(tdir, f)
            cfg = _find_config(tpl)
            if not cfg:
                print(f'[SKIP] {f}（缺少配置，运行 --scan {tpl} 生成）')
                continue
            found = True
            generate(tpl, cfg)
        if not found:
            print('没有找到可处理的模板，请检查模板目录或使用 --template 指定')


if __name__ == '__main__':
    main()
